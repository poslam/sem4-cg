from docx import Document
from docx.shared import Inches

config = {
    "guest": "Петровна Василила Американовна",
    "whose_wedding": "Некрика Кирилла Андреева и Мариновой Ларисы Олеговны",
    "wedding_date": "12 июня 2024 года",
    "wedding_time": "12:00",
    "wedding_place": "ресторан \"У Андрея\" на Светланской 25",
    "guests": [
        {"number": 1, "name": "Иванов Иван Иванович", "relation": "друг жениха",},
        {"number": 2, "name": "Петров Петр Николавевич", "relation": "напарник жениха в доте",},
        {"number": 3, "name": "Сергеевко Лариса Витальевна", "relation": "мастер по ноготочкам невесты",},
        {"number": 4, "name": "Маврения Диана Кирилловна", "relation": "подруга невесты",},
    ]
}

document = Document()

document.add_heading("Приглашение на свадьбу", 0)

p = document.add_paragraph("Приглашаем Вас, ")
p.add_run(config["guest"]).bold = True
p.add_run(", на свадьбу ")
p.add_run(f"{config['whose_wedding']}.").italic = True

p = document.add_paragraph(
    f"Свадьба состоится {config["wedding_date"]} в {config["wedding_time"]} в месте: {config["wedding_place"]}."
)

document.add_heading("Стоит взять с собой:", level=1)

document.add_paragraph("Хорошее настроение", style="List Number")
document.add_paragraph("Счастливую мордашку", style="List Number")
document.add_paragraph("Подарок для брачующихся", style="List Number")

document.add_heading("Список гостей:", level=1)

table = document.add_table(rows=1, cols=3, style="Table Grid")
hdr_cells = table.rows[0].cells
hdr_cells[0].text = "Номер"
hdr_cells[1].text = "ФИО"
hdr_cells[2].text = "Отношение"

for guest in config["guests"]:
    row_cells = table.add_row().cells
    for i in range(3):
        row_cells[i].text = str(list(guest.values())[i])

document.add_picture("./assets/in/sign.png", width=Inches(1.25))
document.add_picture("./assets/in/stamp.png", width=Inches(1.25))

document.save("./assets/out/2.docx")
